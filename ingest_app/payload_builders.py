from __future__ import annotations

import zipfile
from pathlib import Path
from typing import Any

import fitz
from docx import Document
from docx.text.paragraph import Paragraph

try:
    from ingest_app.file_utils import chunk_list
    from ingest_app.text_utils import (
        clean_text, detect_lang, token_count_simple,
        build_structured_json, build_markdown_text, extract_links,
        normalize_text, apply_content_filter,
    )
except ModuleNotFoundError:
    from file_utils import chunk_list
    from text_utils import (
        clean_text, detect_lang, token_count_simple,
        build_structured_json, build_markdown_text, extract_links,
        normalize_text, apply_content_filter,
    )

def build_asset_prefix(file_hash: str) -> str:
    return f"file_{file_hash[:12]}"

def extract_pdf_assets(page: fitz.Page, asset_prefix: str) -> list[dict[str, Any]]:
    data = page.get_text("dict")
    blocks = data.get("blocks", [])
    assets = []
    asset_no = 1

    for block in blocks:
        if block.get("type") == 1 and block.get("image"):
            bbox = block.get("bbox") or []
            assets.append({
                "asset_id": f"{asset_prefix}_p{page.number + 1}_a{asset_no}",
                "asset_type": "image",
                "image_type": "embedded_image",
                "bbox": {
                    "x1": bbox[0] if len(bbox) > 0 else None,
                    "y1": bbox[1] if len(bbox) > 1 else None,
                    "x2": bbox[2] if len(bbox) > 2 else None,
                    "y2": bbox[3] if len(bbox) > 3 else None
                },
                "storage_uri": None,
                "ocr_text": None,
                "caption": None,
                "vision_summary": None,
                "tags": ["embedded_image"],
                "ocr_confidence": None,
                "vision_confidence": None,
                "is_relevant": True
            })
            asset_no += 1
    return assets

def build_pdf_payload(file_path: Path, file_hash: str) -> dict[str, Any]:
    with fitz.open(file_path) as doc:
        asset_prefix = build_asset_prefix(file_hash)
        pages = []
        all_text = []
        total_pdf_pages = len(doc)

        for page in doc:
            text_raw = page.get_text("text", sort=True) or ""
            text_raw = clean_text(text_raw)
            if not text_raw:
                continue
            all_text.append(text_raw)
            assets = extract_pdf_assets(page, asset_prefix)

            pages.append({
                "page_no": page.number + 1,
                "text_raw": text_raw,
                "language": detect_lang(text_raw),
                "char_count": len(text_raw),
                "token_count": token_count_simple(text_raw),
                "page_metadata": {
                    "has_tables": None,
                    "has_images": len(assets) > 0,
                    "has_footnotes": None
                },
                "assets": assets,
                "ocr_confidence": None
            })

        combined_text = "\n\n".join(t for t in all_text if t)
        raw_cleaned_content = combined_text
        structured_json = build_structured_json(file_hash, file_path.name, "pdf", raw_cleaned_content)
        markdown_text = structured_json["markdown_text"]

        return {
            "doc_id": file_hash,
            "file_name": file_path.name,
            "file_path": str(file_path).replace("\\", "/"),
            "file_hash": file_hash,
            "source_type": "pdf",
            "extraction_status": "success",
            "language": structured_json.get("language", "unknown"),
            "page_count": total_pdf_pages,
            "extracted_page_count": len(pages),
            "links": extract_links(raw_cleaned_content),
            "content": raw_cleaned_content,
            "structured_json": structured_json,
            "markdown_text": markdown_text,
            "pages": pages,
        }

def extract_docx_media_count(file_path: Path) -> int:
    try:
        with zipfile.ZipFile(file_path, "r") as zf:
            return sum(1 for name in zf.namelist() if name.startswith("word/media/"))
    except Exception:
        return 0

def build_docx_payload(file_path: Path, file_hash: str, logical_page_paragraphs: int = 20) -> dict[str, Any]:
    doc = Document(file_path)
    asset_prefix = build_asset_prefix(file_hash)

    # Collect paragraph texts and table texts in document order
    content_blocks: list[dict[str, Any]] = []
    has_tables = len(doc.tables) > 0

    # Iterate through document body elements in order to preserve
    # the natural interleaving of paragraphs and tables
    table_idx = 0
    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        if tag == 'p':
            # Paragraph element — extract text
            para = Paragraph(element, doc)
            txt = para.text or ""
            if txt.strip():
                content_blocks.append({"type": "paragraph", "text": txt})
        elif tag == 'tbl':
            # Table element — extract as formatted text
            if table_idx < len(doc.tables):
                table = doc.tables[table_idx]
                table_idx += 1
                rows = []
                for row in table.rows:
                    cells = [(cell.text or "").strip() for cell in row.cells]
                    rows.append(" | ".join(cell for cell in cells if cell))
                table_text = "\n".join(r for r in rows if r)
                table_text = clean_text(table_text)
                if table_text.strip():
                    content_blocks.append({"type": "table", "text": table_text})

    media_count = extract_docx_media_count(file_path)

    # Chunk content blocks into logical pages
    paragraph_texts = [b["text"] for b in content_blocks if b["type"] == "paragraph"]
    logical_pages_raw = chunk_list(paragraph_texts, logical_page_paragraphs) if paragraph_texts else [[]]

    pages = []
    all_text = []
    para_idx = 0  # tracks position in content_blocks

    for idx, para_group in enumerate(logical_pages_raw, start=1):
        # For each logical page, collect paragraphs and any tables that follow them
        page_texts: list[str] = []
        page_has_tables = False

        for para_text in para_group:
            page_texts.append(para_text)
            para_idx += 1
            # Check if any table blocks follow this paragraph in content_blocks
            while para_idx < len(content_blocks) and content_blocks[para_idx]["type"] == "table":
                page_texts.append(content_blocks[para_idx]["text"])
                page_has_tables = True
                para_idx += 1

        text_raw = "\n".join(page_texts).strip()
        text_raw = clean_text(text_raw)
        if not text_raw:
            continue
        all_text.append(text_raw)

        pages.append({
            "page_no": idx,
            "text_raw": text_raw,
            "language": detect_lang(text_raw),
            "char_count": len(text_raw),
            "token_count": token_count_simple(text_raw),
            "page_metadata": {
                "has_tables": page_has_tables or (has_tables if idx == 1 else None),
                "has_images": None,
                "has_footnotes": None
            },
            "assets": []
        })

    # Handle any remaining table blocks not yet captured
    while para_idx < len(content_blocks):
        block = content_blocks[para_idx]
        if block["type"] == "table" and block["text"].strip():
            all_text.append(block["text"])
            pages.append({
                "page_no": len(pages) + 1,
                "text_raw": block["text"],
                "language": detect_lang(block["text"]),
                "char_count": len(block["text"]),
                "token_count": token_count_simple(block["text"]),
                "page_metadata": {
                    "has_tables": True,
                    "has_images": None,
                    "has_footnotes": None
                },
                "assets": []
            })
        para_idx += 1

    combined_text = "\n\n".join(t for t in all_text if t)
    raw_cleaned_content = combined_text
    structured_json = build_structured_json(file_hash, file_path.name, "docx", raw_cleaned_content)
    markdown_text = structured_json["markdown_text"]

    return {
        "doc_id": file_hash,
        "file_name": file_path.name,
        "file_path": str(file_path).replace("\\", "/"),
        "file_hash": file_hash,
        "source_type": "docx",
        "extraction_status": "success",
        "language": structured_json.get("language", "unknown"),
        "page_count": len(pages),
        "extracted_page_count": len(pages),
        "media_count": media_count,
        "links": extract_links(raw_cleaned_content),
        "content": raw_cleaned_content,
        "structured_json": structured_json,
        "markdown_text": markdown_text,
        "pages": pages,
    }

def build_txt_payload(file_path: Path, file_hash: str) -> dict[str, Any]:
    raw = file_path.read_text(encoding="utf-8", errors="ignore")
    raw_cleaned_content = clean_text(raw)
    if not raw_cleaned_content:
        return {
            "doc_id": file_hash,
            "file_name": file_path.name,
            "file_path": str(file_path).replace("\\", "/"),
            "file_hash": file_hash,
            "source_type": "txt",
            "extraction_status": "empty",
            "language": "unknown",
            "page_count": 0,
            "extracted_page_count": 0,
            "links": [],
            "content": "",
            "structured_json": {},
            "markdown_text": "",
            "pages": [],
        }

    structured_json = build_structured_json(file_hash, file_path.name, "txt", raw_cleaned_content)
    markdown_text = structured_json["markdown_text"]

    return {
        "doc_id": file_hash,
        "file_name": file_path.name,
        "file_path": str(file_path).replace("\\", "/"),
        "file_hash": file_hash,
        "source_type": "txt",
        "extraction_status": "success",
        "language": structured_json.get("language", "unknown"),
        "page_count": 1,
        "extracted_page_count": 1,
        "links": extract_links(raw_cleaned_content),
        "content": raw_cleaned_content,
        "structured_json": structured_json,
        "markdown_text": markdown_text,
        "pages": [
            {
                "page_no": 1,
                "text_raw": raw_cleaned_content,
                "language": detect_lang(raw_cleaned_content),
                "char_count": len(raw_cleaned_content),
                "token_count": token_count_simple(raw_cleaned_content),
                "page_metadata": {
                    "has_tables": None,
                    "has_images": None,
                    "has_footnotes": None
                },
                "assets": []
            }
        ],
    }